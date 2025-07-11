import os
import io
import base64
import logging
from typing import Optional

import streamlit as st
from dotenv import load_dotenv

import openai
import pandas as pd

from prompts import CR_PROMPT, MC_PROMPT, MS_PROMPT, EBSR_PROMPT, TE_PROMPT

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global PDF reader
try:
    import pypdf  # type: ignore
    PDF_READER_CLASS = pypdf.PdfReader
    logger.info("Successfully imported pypdf")
except ImportError:
    logger.warning("Failed to import pypdf, trying PyPDF2")
    try:
        import PyPDF2  # type: ignore
        PDF_READER_CLASS = PyPDF2.PdfReader
        logger.info("Successfully imported PyPDF2")
    except ImportError as e:
        logger.error(f"Failed to import PDF libraries: {e}")
        st.error("PDF processing libraries not available. Please contact support.")
        PDF_READER_CLASS = None

# Load environment variables
load_dotenv()

def clean_extracted_text(text: str) -> str:
    """Clean and format extracted PDF text."""
    if not text:
        return ""
    text = text.replace('\n\n', '\n')
    return text.strip()

def extract_pdf_content(pdf_path: str) -> Optional[str]:
    """Extract and process text content from a PDF file."""
    if not PDF_READER_CLASS:
        st.error("PDF processing is not available")
        return None
    logger.info(f"Starting PDF extraction from: {pdf_path}")
    try:
        if not os.path.exists(pdf_path):
            logger.error(f"PDF file not found: {pdf_path}")
            return None
        with open(pdf_path, 'rb') as file:
            reader = PDF_READER_CLASS(file)
            text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
            return clean_extracted_text(" ".join(text))
    except Exception as e:
        logger.error(f"Error in PDF extraction: {str(e)}")
        return None

def load_references(grade: str) -> tuple[str, str, str]:
    """Load 3D NGSS, and DOK Levels reference texts."""
    try:
        ngss_text = extract_pdf_content("reference_materials/3D NGSS.pdf") or ""
    except Exception as e:
        logger.error(f"Error extracting 3D NGSS.pdf: {e}")
        ngss_text = f"[ERROR: Could not extract 3D NGSS.pdf: {e}]"
    try:
        dok_text = extract_pdf_content("reference_materials/DOK Levels.pdf") or ""
    except Exception as e:
        logger.error(f"Error extracting DOK Levels.pdf: {e}")
        dok_text = f"[ERROR: Could not extract DOK Levels.pdf: {e}]"
    return ngss_text, dok_text

def get_response(
    grade: str,
    item_type: str,
    num_items: int,
    standards: str,
    will_do: str,
) -> str:
    """Generate assessment content using the OpenAI GPT-4o model with updated prompt."""
    ngss_ref, dok_ref = load_references(grade)
    if item_type == "Constructed Response":
        prompt_template = CR_PROMPT
    elif item_type == "Multiple Choice":
        prompt_template = MC_PROMPT
    elif item_type == "Multiple Select":
        prompt_template = MS_PROMPT
    elif item_type == "Evidence-Based":
        prompt_template = EBSR_PROMPT
    elif item_type == "Technology Enhanced":
        prompt_template = TE_PROMPT
    else:
        prompt_template = ''

    prompt = prompt_template.format(
        grade=grade,
        num_items=num_items,
        standard=standards,
        will_do=will_do,
        ngss_ref=ngss_ref,
        dok_ref=dok_ref,
    )
    response = openai.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "You are a science assessment writer who exactly replicates official state assessment style and format. Output ONLY the questions and their solutions."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4000,
        temperature=0.7
    )
    return response.choices[0].message.content


# Streamlit UI
st.title("Science Assessment Generator")
st.subheader("Generate science assessment items by type and count")
st.markdown(
    """
Refer to the appropriate tab in this [spreadsheet](https://docs.google.com/spreadsheets/d/1lXPIsCrwuEH3yAaj5xR56-FloC-pXU03L7gdh4Qe7bg/edit?usp=sharing) and copy into the corresponding textboxes.
    """
)
standards = pd.read_csv("reference_materials/standards.csv")
# Ensure grade_unit is string for matching
if 'grade_unit' in standards.columns:
    standards['grade_unit'] = standards['grade_unit'].astype(str)
# UI Row 1: Grade, Unit, Item Type, Number of Items
col1, col2, col3, col4 = st.columns(4)
with col1:
    grade = st.selectbox(
        "Grade Level:",
        [f"{i}" for i in range(6, 9)]
    )
with col2:
    unit = st.selectbox(
        "Unit:",
        [f"{i}" for i in range(1, 7)]
    )
    grade_unit = f"{grade}.{unit}"
    
with col3:
    item_type = st.selectbox(
        "Item Type:",
        ["Multiple Choice", "Multiple Select", "Technology Enhanced", "Cluster", "Evidence-Based", "Constructed Response"]
    )
with col4:
    # Limit max to 3 for Cluster, EBSR, and CR; else 3-10
    if item_type == "Cluster":
        num_items = st.selectbox(
            "Number of Clusters:",
            list(range(1, 4))
        )
    elif item_type == "Evidence-Based":
        num_items = st.selectbox(
            "Number of EBSR Sets:",
            list(range(1, 4))
        )
    elif item_type == "Constructed Response":
        num_items = st.selectbox(
            "Number of CR Items:",
            list(range(1, 4))
        )
    else:
        num_items = st.selectbox(
            "Number of Items:",
            list(range(3, 11))
        )


# UI Row 2: Standards (selection) & Figure Out
r1c1, r1c2 = st.columns(2)
with r1c1:
    
    # Filter standards DataFrame for the selected grade_unit
    filtered_standards = standards[standards['grade_unit'] == grade_unit] if 'grade_unit' in standards.columns else standards
    standard_options = filtered_standards['Standards'].tolist() if 'Standards' in filtered_standards.columns else []

    if standard_options:
        standards_selected = st.selectbox("Select Standards:", options=standard_options)
        standards = '\n'.join(standards_selected)
    else:
        standards = st.text_area("Standards (no standards found for this grade/unit):", height=150)
with r1c2:
    will_dos = pd.read_csv("reference_materials/will_do.csv", encoding='latin-1')
    # convert grade_unit to string
    will_dos['grade_unit'] = will_dos['grade_unit'].astype(str)
    # filter will_dos for the selected grade_unit
    filtered_will_dos = will_dos[will_dos['grade_unit'] == grade_unit] if 'grade_unit' in will_dos.columns else will_dos
    will_do = filtered_will_dos.iloc[0]['will_do'] if 'will_do' in filtered_will_dos.columns else ''
    st.text_area("What Students Will Do:", value=will_do, height=150)



if 'last_params' not in st.session_state:
    st.session_state['last_params'] = None
if 'last_results' not in st.session_state:
    st.session_state['last_results'] = None

params = (grade, unit, item_type, num_items, standards, will_do)

if st.button("Generate Assessment"):
    st.session_state['last_params'] = params
    st.session_state['last_results'] = None

if st.session_state.get('last_params') == params and st.session_state.get('last_results') is not None:
    all_results, file_name, buffer = st.session_state['last_results']
    st.download_button(
        label="Download as Word (.docx)",
        data=buffer,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.markdown(all_results, unsafe_allow_html=True)
else:
    if st.session_state.get('last_params') == params:
        import re
        from docx import Document
        all_results = ""
        doc = Document()
        try:
            status_placeholder = st.empty()
            if item_type == "Cluster":
                status_placeholder.info("Generating clusters with 2 MC, 2 MS, 2 TE, and 2 random (MC/MS/TE) items each...")
                for cluster_num in range(1, num_items + 1):
                    # 2 MC
                    with st.spinner(f"Generating Cluster {cluster_num} Items 1–2 (MC)..."):
                        batch_mc = get_response(
                            grade, unit, "MC", 2, standards, will_do, "", item_start=1, item_end=2, batch_num=1, te_type=""
                        )
                        all_results += batch_mc + "\n\n"
                        for line in batch_mc.split("\n"):
                            doc.add_paragraph(line)
                    # 2 MS
                    with st.spinner(f"Generating Cluster {cluster_num} Items 3–4 (MS)..."):
                        batch_ms = get_response(
                            grade, unit, "MS", 2, standards, will_do, "", item_start=3, item_end=4, batch_num=2, te_type=""
                        )
                        all_results += batch_ms + "\n\n"
                        for line in batch_ms.split("\n"):
                            doc.add_paragraph(line)
                    # 2 TE (rotate subtype for variety, e.g., Drag-and-Drop and Hot-Spot)
                    te_types = [" - Drag-and-Drop", " - Hot-Spot"]
                    for i, te_subtype in enumerate(te_types):
                        with st.spinner(f"Generating Cluster {cluster_num} Item {5+i} (TE{te_subtype})..."):
                            batch_te = get_response(
                                grade, unit, "TE", 1, standards, will_do, "", item_start=5+i, item_end=5+i, batch_num=3+i, te_type=te_subtype
                            )
                            all_results += batch_te + "\n\n"
                            for line in batch_te.split("\n"):
                                doc.add_paragraph(line)
                    # 2 random (MC, MS, or TE)
                    import random
                    random_types = ["MC", "MS", "TE"]
                    for i in range(7, 9):
                        rand_type = random.choice(random_types)
                        te_type_val = ""
                        if rand_type == "TE":
                            te_type_val = random.choice([" - Drag-and-Drop", " - Hot-Spot", " - Inline-Choice", " - Graphing"])
                        with st.spinner(f"Generating Cluster {cluster_num} Item {i} (Random: {rand_type}{te_type_val})..."):
                            batch_rand = get_response(
                                grade, unit, rand_type, 1, standards, will_do, "", item_start=i, item_end=i, batch_num=5+i, te_type=te_type_val
                            )
                            all_results += batch_rand + "\n\n"
                            for line in batch_rand.split("\n"):
                                doc.add_paragraph(line)
                file_name = f"{grade} {unit} {item_type} Items.docx"
            elif item_type == "Evidence-Based":
                status_placeholder.info("Generating EBSR item sets...")
                for ebsr_num in range(1, num_items + 1):
                    with st.spinner(f"Generating EBSR Set {ebsr_num}..."):
                        ebsr_result = get_response(
                            grade, item_type, 1, standards, will_do
                        )
                        all_results += ebsr_result + "\n\n"
                        for line in ebsr_result.split("\n"):
                            doc.add_paragraph(line)
                file_name = f"{grade} {unit} {item_type} Sets.docx"
            elif item_type in ["Constructed Response", "Multiple Choice", "Multiple Select", "Technology Enhanced"]:
                status_placeholder.info(f"Generating {item_type} items...")
                for cr_num in range(1, num_items + 1):
                    with st.spinner(f"Generating {item_type} Item {cr_num}..."):
                        cr_result = get_response(
                            grade, item_type, 1, standards, will_do
                        )
                        all_results += cr_result + "\n\n"
                        for line in cr_result.split("\n"):
                            doc.add_paragraph(line)
                file_name = f"{grade} {unit} {item_type} Items.docx"
            else:
                status_placeholder.info("Generating assessment items...")
                batch_size = 10
                total = int(num_items)
                current = 1
                while current <= total:
                    end = min(current + batch_size - 1, total)
                    batch_result = get_response(
                        grade, unit, item_type, end - current + 1, standards, will_do, ""
                    )
                    all_results += batch_result + "\n\n"
                    for line in batch_result.split("\n"):
                        doc.add_paragraph(line)
                    matches = list(re.finditer(r"Item (\d+):", batch_result))
                    if matches:
                        last_item = int(matches[-1].group(1))
                        current = last_item + 1
                    else:
                        current = end + 1
                file_name = f"{grade} {unit} {item_type} Item Types.docx"
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            status_placeholder.success("Item generation complete!")
            st.session_state['last_results'] = (all_results, file_name, buffer)
            st.download_button(
                label="Download as Word (.docx)",
                data=buffer,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.markdown(all_results, unsafe_allow_html=True)
        except Exception as e:
            import traceback
            tb = traceback.format_exc()
            st.error(f"An error occurred: {e}\n\nTraceback:\n{tb}")
            logger.error(f"Error generating assessment: {e}\n{tb}")
    elif not all([grade, unit, item_type, num_items, standards, will_do]):
        st.warning("Please fill in all fields.")
